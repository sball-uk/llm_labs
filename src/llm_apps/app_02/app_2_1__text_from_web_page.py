"""This module copies the main text from a web page to a Word file.
See app_02/README.md
"""

from datetime import datetime
import requests

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document

# Specify input and output directories and filenames
DIR_INPUT = "../data/app_02__skills_toolbox/2.1.1__input_sources/"
FNAME_INPUT = "2.1.1__input_sources.csv"
DIR_OUTPUT = "../data/app_02__skills_toolbox/2.1.9__output_downloaded/"


if __name__ == "__main__":

    # Import csv, filter and create dictionary
    df = pd.read_csv(DIR_INPUT + FNAME_INPUT)
    df_fil = df[(df["to_get"] == 1) & df["tb_url"].notna()]
    dict_urls_to_get = dict(zip(df_fil["tb_id"], df_fil["tb_url"]))

    # Loop over all ids in the dict
    if len(dict_urls_to_get) == 0:
        print("No urls to download")
    else:
        list_status_code = []
        for tb_id, tb_url in dict_urls_to_get.items():
            print("")
            print(f"Key: {tb_id}, Value: {tb_url}")

            # Create a new Word document
            doc = Document()

            # Add the hyperlink
            doc.add_paragraph(tb_url)

            # When using requests, some websites block requests that do not resemble typical browser requests.
            # To mimic a browser we add a 'User-Agent' header, which should convert some 403 status codes to 200.
            headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 \
                                        (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36"}

            # Get data from the website
            try:
                response = requests.get(tb_url, headers=headers, timeout=30)
                response_status_code = response.status_code

            except Exception as e:
                print(f"An error occurred: {e}")
                response = 'N/A'
                response_status_code = 88888
                doc.add_paragraph("Error: issue with source url, no response")

            # Capture response status codes as a list of dicts
            print("Response:", response_status_code)
            dict_status_code = {
                'tb_id': tb_id,
                'status_code': response_status_code}
            list_status_code.append(dict_status_code)

            # Extract 'main' text for successful responses, i.e. status == 200
            if response_status_code == 200:

                # Extract title and paragraph elements from the webpage
                soup = BeautifulSoup(response.text, "html.parser")

                # Find the desired tags (p, title, etc.)
                all_text = soup.find_all(["p", "title", "h1", "h2", "h3"])

                # Iterate through tags and convert them to Word format
                # p: Paragraphs are added as plain text using doc.add_paragraph().
                # title and h1, h2, h3: Headings are added with doc.add_heading() at different levels.
                for tag in all_text:
                    if tag.name == 'p':  # For paragraphs
                        try:
                            doc.add_paragraph(tag.get_text())
                        except Exception as e:
                            print(f"An error occurred: {e}")

                    elif tag.name == 'title':  # For titles use Heading Level 1
                        try:
                            doc.add_heading(tag.get_text(), level=1)
                        except Exception as e:
                            print(f"An error occurred: {e}")

                    elif tag.name == 'h1':  # For h1 headings
                        try:
                            doc.add_heading(tag.get_text(), level=1)
                        except Exception as e:
                            print(f"An error occurred: {e}")

                    elif tag.name == 'h2':  # For h2 headings
                        try:
                            doc.add_heading(tag.get_text(), level=2)
                        except Exception as e:
                            print(f"An error occurred: {e}")

                    elif tag.name == 'h3':  # For h3 headings
                        try:
                            doc.add_heading(tag.get_text(), level=3)
                        except Exception as e:
                            print(f"An error occurred: {e}")

            # Save
            doc.save(DIR_OUTPUT + tb_id + "_auto_" + datetime.now().strftime("%Y-%m-%d") + ".docx")

            del response
            del response_status_code

        # save status_codes to logs
        df_status_code = pd.DataFrame(list_status_code)
        df_status_code.to_csv("../logs/log_2_1__status_codes.txt", mode="a", header=False, index=False)
