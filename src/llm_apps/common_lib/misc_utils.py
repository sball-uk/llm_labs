"""This module provides common functions for miscellaneous utilities.
"""

import os

from dotenv import load_dotenv
from docx import Document


def load_dotenv_all():
    """Load the file containing the user's local environment variables, e.g. API keys.

    This file is excluded from git by .gitignore.

    Typical usage:
        clm.load_dotenv_all()
    """
    load_dotenv("../.env", override=True)


def list_docx_in_directory(dir_docx: str):
    """List all docx files in the specified directory.

    Args:
        directory (str): The directory to list files from; defaults to the current working directory.

    Returns:
        list: A list of docx files in specified directory.

    Raises:
        Exception: Any error.

    Typical usage:
        list_docs_to_review = clm.list_docx_in_directory(DIR_INPUT_DOCX)
    """

    try:
        files = [f for f in os.listdir(dir_docx) if ".docx" in f]
    except Exception as e:
        print(f"An error occurred: {e}")

    return files


def load_docx_to_str(path_docx: str):
    """Load the text from a docx file.

    Args:
        path_docx (str): The path of the docx file.

    Returns:
        str: The first paragraph of the file; the convention used is for this to be the url to the source material.
        str: All text in the document.

    Typical usage:
        doc_first_para, doc_all_text = clm.load_docx_to_str(DIR_INPUT_DOCX + FNAME_INPUT_DOCX)
    """

    # Load the .docx file
    doc = Document(path_docx)

    # Get first paragraph, e.g. document url
    first_paragraph = doc.paragraphs[0].text

    # Iterate through each paragraph and concatenate the text
    all_text = ""
    for paragraph in doc.paragraphs:
        all_text += paragraph.text + "\n"

    return first_paragraph, all_text
